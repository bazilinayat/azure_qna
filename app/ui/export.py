"""Export a chat transcript to .docx or .pdf.

Both return raw bytes so Streamlit can hand them straight to a download button
without touching the filesystem — which matters, because the whole point of the
session-only design is that nothing is written to disk.
"""

from datetime import datetime
from io import BytesIO
import re

# reportlab's built-in fonts are Latin-1, and LLM output is full of curly quotes,
# em dashes and arrows that are not in it. Embedding a Unicode TTF would mean
# shipping a font file, so the handful of characters that actually show up get
# folded to ASCII instead.
_ASCII_REPLACEMENTS = {
    "‘": "'", "’": "'",
    "“": '"', "”": '"',
    "–": "-", "—": "-",
    "…": "...", "•": "-",
    " ": " ", "→": "->",
    "×": "x", "≥": ">=", "≤": "<=",
}


def to_ascii(text: str) -> str:
    for source, target in _ASCII_REPLACEMENTS.items():
        text = text.replace(source, target)

    return text.encode("ascii", "replace").decode("ascii")


def _strip_markdown(text: str) -> str:
    """Flatten the markdown the model emits into plain text.

    Neither exporter renders markdown, so leaving the syntax in means readers
    see literal ** and ### in their saved document.
    """

    text = re.sub(r"```[a-zA-Z]*\n?", "", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)

    return text.strip()


def _timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def default_filename(extension: str) -> str:
    return f"azurementor-chat-{datetime.now():%Y%m%d-%H%M}.{extension}"


# --------------------------------------------------
# Word
# --------------------------------------------------

def to_docx(messages: list[dict]) -> bytes:
    """Render the transcript as a .docx file.

    `messages` is the Streamlit history: dicts with `role`, `content` and, for
    assistant turns, an optional `sources` list.
    """

    from docx import Document
    from docx.shared import Pt

    document = Document()

    document.add_heading("AzureMentor conversation", level=0)

    subtitle = document.add_paragraph(f"Exported {_timestamp()}")
    subtitle.runs[0].font.size = Pt(9)

    document.add_paragraph(
        "Answers are generated from Microsoft Azure documentation and may "
        "contain errors. Verify anything important against the linked sources."
    ).runs[0].font.size = Pt(9)

    for message in messages:

        if message["role"] == "user":
            document.add_heading(message["content"], level=2)
            continue

        for block in _strip_markdown(message["content"]).split("\n\n"):
            if block.strip():
                document.add_paragraph(block.strip())

        sources = message.get("sources") or []

        if sources:
            document.add_paragraph("Sources").runs[0].bold = True

            for index, source in enumerate(sources, start=1):
                paragraph = document.add_paragraph(
                    f"[{index}] {source['title']}\n{source['url']}",
                    style="List Bullet",
                )
                paragraph.runs[0].font.size = Pt(9)

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()


# --------------------------------------------------
# PDF
# --------------------------------------------------

def to_pdf(messages: list[dict]) -> bytes:
    """Render the transcript as a PDF."""

    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="AzureMentor conversation",
    )

    styles = getSampleStyleSheet()

    question_style = ParagraphStyle(
        "Question",
        parent=styles["Heading2"],
        spaceBefore=14,
        spaceAfter=6,
    )

    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        alignment=TA_LEFT,
        leading=14,
        spaceAfter=6,
    )

    small_style = ParagraphStyle(
        "Small",
        parent=styles["BodyText"],
        fontSize=8,
        leading=10,
        textColor="#555555",
    )

    story = [
        Paragraph("AzureMentor conversation", styles["Title"]),
        Paragraph(f"Exported {_timestamp()}", small_style),
        Paragraph(
            "Answers are generated from Microsoft Azure documentation and may "
            "contain errors. Verify anything important against the linked "
            "sources.",
            small_style,
        ),
        Spacer(1, 12),
    ]

    for message in messages:

        if message["role"] == "user":
            story.append(
                Paragraph(_escape(message["content"]), question_style)
            )
            continue

        for block in _strip_markdown(message["content"]).split("\n\n"):
            if block.strip():
                story.append(
                    Paragraph(
                        _escape(block.strip()).replace("\n", "<br/>"),
                        body_style,
                    )
                )

        sources = message.get("sources") or []

        if sources:
            story.append(Spacer(1, 4))
            story.append(Paragraph("<b>Sources</b>", small_style))

            for index, source in enumerate(sources, start=1):
                story.append(
                    Paragraph(
                        f"[{index}] {_escape(source['title'])}<br/>"
                        f"{_escape(source['url'])}",
                        small_style,
                    )
                )

        story.append(Spacer(1, 8))

    document.build(story)

    return buffer.getvalue()


def _escape(text: str) -> str:
    """Escape for reportlab's mini-HTML, after folding to ASCII."""

    return (
        to_ascii(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
